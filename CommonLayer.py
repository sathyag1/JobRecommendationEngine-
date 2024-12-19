import DataLayer as DL
import BusinessLayer as BL
import numpy as np
import spacy
nlp = spacy.load('en_core_web_sm')
from sentence_transformers import SentenceTransformer,util
model = SentenceTransformer('all-MiniLM-L6-v2')
import fitz
import pandas as pd
import json
import os
from dotenv import load_dotenv
from elasticapm.contrib.flask import ElasticAPM
from elasticapm.handlers.logging import LoggingHandler

load_dotenv()



import docx
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.document import Document
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph

from azure.storage.blob import BlobServiceClient
from azure.storage.blob import BlobServiceClient, BlobClient, ContainerClient

from geopy.geocoders import Nominatim
from geopy import distance
geocoder = Nominatim(user_agent="I know python")

AZURE_BLOB_CONTAINERNAME = os.getenv('AZURE_BLOB_CONTAINERNAME')
AZURE_BLOB_CONNECTION_STRING = os.getenv('AZURE_BLOB_CONNECTION_STRING')


#blob Container
CONTAINERNAME= AZURE_BLOB_CONTAINERNAME

CONNECTION_STRING=AZURE_BLOB_CONNECTION_STRING

def convert_array(x):
    n12 = np.squeeze(np.array(x))
    return n12

def paginate_dataframe(job_recommendation_relevancy, page_size, page_num):

    if page_size is None:

        return None

    offset = page_size*(page_num-1)

    return job_recommendation_relevancy[offset:offset + page_size]

#Function for vector padding (address dimension mismatch) and calculating cosine similarity score
def cosine_sim(a, b): 
    a = np.array(a)
    size=""
    arr=""
    if a.shape == b.shape:
        cos_sim = np.dot(a, b)/(np.linalg.norm(a)*np.linalg.norm(b))
        return cos_sim
    else:
         

        if a.size > b.size:
                size = a.size - b.size
                arr = np.zeros(size)
                b = np.concatenate((b, arr))               
                cos_sim = np.dot(a, b)/(np.linalg.norm(a)*np.linalg.norm(b))

                return cos_sim
        else:   
                size = b.size - a.size
                arr = np.zeros(size)
                a = np.concatenate((a, arr))
                cos_sim = np.dot(a, b)/(np.linalg.norm(a)*np.linalg.norm(b))
                return cos_sim

"##### Function for covert JobDescription into Vectors"
def convert_sentancevec(x):
       try:
             jd=clean_text(str(x))
             sentence_vec=model.encode(jd),
       except ValueError: 
            sentence_vec="0"
       return sentence_vec

"##### Function for preprocessing of text"
def clean_text(text):
    doc = nlp(text)
    keywords = [str(token.lemma_).lower() for token in doc if token.pos_ not in ['PRON', 'PUNCT', 'SPACE'] and not token.is_stop and len(token.text)>1]
    clean_text = ' '.join(keywords)
    return clean_text  

"##### Function for calculated Matching score based on weightage of parameter"
def matchingscore(row,maximum_salary,maximum_experience,qualification_encoded,skill_vec,preferredjobs,place,userjourney):  
    
    skillscore = 0
    experience_score=0
    education_score=0
    salary_score = 0
    location_score=0
    skill_score=0
    if userjourney == "1":
        salaryweightage=20
        experienceweightage=15
        educationweightage=15
        skillsweightage=25
        locationweightage=25
    else:
        salaryweightage=20
        experienceweightage=15
        educationweightage=15
        profileweightage=25
        locationweightage=25
    def isNaN(num):
        return num!= num    
    if row['MAXIMUM_SALARY'] is not None:
        #isnan=np.isnan(row['MAXIMUM_SALARY'])
        isnan=isNaN(row['MAXIMUM_SALARY'])
        if isnan is False and row['MAXIMUM_SALARY'] !="" and row['MAXIMUM_SALARY'] !=0:
            
            if  row['MAXIMUM_SALARY'] >= maximum_salary:
                #salary_score=salaryweightage
                #add the new logic here
                salary_score=salaryweightage-((abs(maximum_salary-row['MAXIMUM_SALARY'])/row['MAXIMUM_SALARY'])*salaryweightage)
            elif row['MAXIMUM_SALARY'] < maximum_salary:
                salary_score=row['MAXIMUM_SALARY']/maximum_salary*salaryweightage
                
        
        elif isNaN(row['MINIMUM_SALARY']) is False and row['MINIMUM_SALARY'] !="" and row['MINIMUM_SALARY'] !=0:
            
            if  row['MINIMUM_SALARY'] >= maximum_salary:
                #salary_score=salaryweightage
                #add the new logic here
                salary_score=salaryweightage-((abs(maximum_salary-row['MINIMUM_SALARY'])/row['MINIMUM_SALARY'])*salaryweightage)
            #elif row['MINIMUM_SALARY'] < maximum_salary:
            else:
                salary_score=row['MINIMUM_SALARY']/maximum_salary*salaryweightage 
                # print(salary_score)   
    #print("salary input",maximum_salary,row['MINIMUM_SALARY'],",",row['MAXIMUM_SALARY'],"slary score",salary_score)
    ########new exp score logic
    
    if maximum_experience==0 and maximum_experience!="":
        #isnan=np.isnan(row['MAXIMUM_EXPERIENCE'])
        isnan=isNaN(row['MAXIMUM_EXPERIENCE'])
        if row['MAXIMUM_EXPERIENCE'] is not None and isnan is False:
            if row['MINIMUM_EXPERIENCE']==0 and row['MAXIMUM_EXPERIENCE'] ==0:
                experience_score=experienceweightage
            else:
                int_max_exp=1
                if  row['MAXIMUM_EXPERIENCE']!=0:
                    experience_score=(int_max_exp/row['MAXIMUM_EXPERIENCE'])*15
                elif row['MAXIMUM_EXPERIENCE']==0 and row['MINIMUM_EXPERIENCE']!=0:
                    experience_score=(int_max_exp/row['MINIMUM_EXPERIENCE'])*15
        else:
            #isnan=np.isnan(row['MINIMUM_EXPERIENCE'])
            isnan=isNaN(row['MINIMUM_EXPERIENCE'])
            if row['MINIMUM_EXPERIENCE'] is not None and isnan is False:
                if row['MINIMUM_EXPERIENCE']==0:
                    experience_score=experienceweightage
                else:
                    int_max_exp=1
                    experience_score=(int_max_exp/row['MINIMUM_EXPERIENCE'])*15
    else:
        #isnan=np.isnan(row['MAXIMUM_EXPERIENCE'])
        isnan=isNaN(row['MAXIMUM_EXPERIENCE'])
        if row['MAXIMUM_EXPERIENCE'] is not None and isnan is False:
            if  row['MAXIMUM_EXPERIENCE'] !="" and row['MAXIMUM_EXPERIENCE'] !=0:
                if row['MINIMUM_EXPERIENCE'] >= maximum_experience:
                    experience_score=(maximum_experience/row['MAXIMUM_EXPERIENCE'])*15
        else:
            #isnan=np.isnan(row['MINIMUM_EXPERIENCE'])
            isnan=isNaN(row['MINIMUM_EXPERIENCE'])
            if row['MINIMUM_EXPERIENCE'] is not None and isnan is False:
                if  row['MINIMUM_EXPERIENCE'] !="" and row['MINIMUM_EXPERIENCE'] !=0:
                    if row['MINIMUM_EXPERIENCE'] >= maximum_experience:
                          experience_score=(maximum_experience/row['MINIMUM_EXPERIENCE'])*15
    #print("for experience input",maximum_experience,"row max and min",row['MAXIMUM_EXPERIENCE'],row['MINIMUM_EXPERIENCE'],"experience_score",experience_score)
    ####newLogic
    
    ##########
    

    
    
    if row['QUALIFICATION_ENCODED'] is not None: 
        if  row['QUALIFICATION_ENCODED'] == 0 or row['QUALIFICATION_ENCODED'] >= qualification_encoded:
            education_score=educationweightage
        elif row['QUALIFICATION_ENCODED'] < qualification_encoded:
            education_score=row['QUALIFICATION_ENCODED']/qualification_encoded*educationweightage
        

    if userjourney == "1":
        if row['SKILL_REQUIRED'] is not None and len(row['SKILL_REQUIRED'])!=0:
            skill=row['SKILLS_VEC']
            skill=convert_array(skill)
            cosinescore = cosine_sim(skill_vec,skill)
            skill_score=(cosinescore/100*skillsweightage)*100
        
    else:
        # print('profilebasedmatch')
        jdvec=row['SENTENCE_VEC']
        jdvec=convert_array(jdvec)
        skillscore = cosine_sim(skill_vec,jdvec)
        skill_score=(skillscore/100*profileweightage)*100  

    if  row['DISTANCE'] <= 20:
        location_score=locationweightage
    elif row['DISTANCE'] <= 40:
        location_score=locationweightage-5
    elif row['DISTANCE'] <= 60:
        location_score=locationweightage-10
    elif row['DISTANCE'] <= 80:
        location_score=locationweightage-15
    elif row['DISTANCE'] <= 100:
        location_score=locationweightage-20
    else:
        location_score=locationweightage-25

    final_score=salary_score+experience_score+education_score+skill_score+location_score
    
    if final_score > 100:
        final_score = 100
    # print(salary_score,experience_score,education_score,skill_score,location_score,final_score)
    return round(final_score)

    

def getqualification_master_encoding(x):
    #generating qualification mapping for correct filtering 
    try:           
            education_mapping  = {
                                "12th": 2,                    
                                "Intermediate" : 3,                               
                                "OTHER" : 0,
                                "10th Pass and below" : 1,
                                "Class 10th" : 1,
                                "10th Pass and above" : 1,
                                "10th pass" : 1,
                                "Below 10th" : 1,
                                "12th Pass and below": 2,
                                "12th pass" : 2,
                                "Class 12th" : 2,
                                "12th Pass and above" : 2,
                                "Diploma/ITI" : 2,
                                "Under Graduate" : 3,
                                "Graduate and above" : 3,
                                "BE" : 3,
                                "Graduate" : 3,
                                "Graduation" : 3,
                                "B.Tech" : 3,
                                "B.A":3,
                                "B.Com":3,
                                "M.Tech":4,
                                "Post Graduate" : 4,
                                "Post Graduate and above" : 4,
                                  '10th': 1,
                                  '11th': 1,
                                  '3D Animation & VFX': 3,
                                  '5th': 1,
                                  '8th': 1,
                                  'Aeronautical Engineering': 3,
                                  'Automation and Robotics': 3,
                                  'Automobile Engineering': 3,
                                  'Aviation Courses': 3,
                                  'B Pharm': 3,
                                  'B Pharm (Pharmaceutical Chemistry)': 3,
                                  'B Pharm (Pharmaceutics)': 3,
                                  'B Pharma (Bachelor of Pharmacy)': 3,
                                  'B.Com- Bachelor of Commerce': 3,
                                  'B.Com (Hons.)': 3,
                                  'B.Com in Computer Applications': 3,
                                  'B.Ed': 3,
                                  'B.Sc': 3,
                                  'B.Sc- Interior Design': 3,
                                  'B.Sc Audiology, Bachelor of Audiology & Speech-Language Pathology (BASLP)': 3,
                                  'B.Sc Dialysis Therapy Technology': 3,
                                  'B.Sc Health Information Management': 3,
                                  'B.Sc in Medical Laboratory Technology': 3,
                                  'B.Sc in Optometry': 3,
                                  'B.Sc Medical Imaging Technology, ': 3,
                                  'B.Sc MRIT': 3,
                                  'B.Sc Nutrition & Dietetics': 3,
                                  'B.Sc Operation Theatre Technology(OTT)': 3,
                                  'B.Sc Radiotherapy Technology(RTT)': 3,
                                  'B.Sc.- Hospitality and Hotel Administration': 3,
                                  'B.Sc. (Hons) Economics with Banking': 3,
                                  'B.Sc. (Hons) in Money, Banking and Finance': 3,
                                  'B.Sc. in Banking and Finance': 3,
                                  'BA (Hons.) in Economics': 3,
                                  'BA in Banking and Finance': 3,
                                  'BA in Banking and Financial Planning': 3,
                                  'BA in History': 3,
                                  'BA in International Finance and Banking': 3,
                                  'Bachelor in Foreign Language': 3,
                                  'Bachelor of Accounting and Finance': 3,
                                  'Bachelor of Ayurvedic Pharmacy': 3,
                                  'Bachelor of Business (Banking and Finance)': 3,
                                  'Bachelor of Business (Banking)': 3,
                                  'Bachelor of Business Administration - International Business (BBA - IB)': 3,
                                  'Bachelor of Business and Commerce (Banking and Finance)': 3,
                                  'Bachelor of Commerce': 3,
                                  'Bachelor of Commerce and Banking Insurance': 3,
                                  'Bachelor of Commerce in Financial Market': 3,
                                  'Bachelor of Design (B. Design)': 3,
                                  'Bachelor of Design in Accessory Design, ': 3,
                                  'Bachelor of Design in Ceramic Design, ': 3,
                                  'Bachelor of Design in fashion Design, ': 3,
                                  'Bachelor of Design in Graphic Design, ': 3,
                                  'Bachelor of Design in Industrial Design, ': 3,
                          'Bachelor of Design in Jewellery Design': 3,
                          'Bachelor of Design in Leather Design, ': 3,
                          'Bachelor of Economics': 3,
                          'Bachelor of Education (B.Ed)': 3,
                          'Bachelor of Elementary Education (B.El.Ed)': 3,
                          'Bachelor of Engg': 3,
                          'Bachelor of Fine Arts (BFA)': 3,
                          'Bachelor of Performing Arts': 3,
                          'Bachelor of Performing Arts (BPA)': 3,
                          'Bachelor of Pharmacy (Hons)': 3,
                          'Bachelor of Pharmacy in Pharmacognosy': 3,
                          'Bachelor of Physical Education (B.P.Ed)': 3,
                          'Bachelor of Visual Arts (BVA)': 3,
                          'BBA- Bachelor of Business Administration': 3,
                          'BBA (Hons) Finance and Banking': 3,
                          'BBA Agriculture Management': 3,
                          'BBS- Bachelor of Business Studies': 3,
                          'BCA': 3,
                          'BCom in Banking': 3,
                          'BE/BTech in Agricultural Engineering': 3,
                          'BEM- Bachelor of Event Management': 3,
                          'BFA- Bachelor of Fine Arts': 3,
                          'BFD- Bachelor of Fashion Designing': 3,
                          'Biotechnology Engineering': 3,
                          'BJMC- Bachelor of Journalism and Mass Communication': 3,
                          'Bleaching and Dyeing Calico Print': 3,
                          'BMS- Bachelor of Management Science': 3,
                          'BSc Agriculture': 3,
                          'BSc Agriculture and Food Business': 3,
                          'BSc Home Science': 3,
                          'BSc in Agriculture Economics and Farm Management': 3,
                          'BSc in Animal Husbandry': 3,
                          'BSc in Crop Physiology': 3,
                          'BSc in Fisheries': 3,
                          'BSc in Forestry': 3,
                          'BSc in Genetic Plant Breeding': 3,
                          'BSc in Horticulture': 3,
                          'BSc Radiography, ': 3,
                          'BSc Soil and water management': 3,
                          'BSW- Bachelor of Social Work': 3,
                          'BTTM- Bachelor of Travel and Tourism Management': 3,
                          'CA- Chartered Accountancy': 3,
                          'Ceramic Engineering': 3,
                          'Chemical Engineering': 3,
                          'Civil Engineering': 3,
                          'Commercial Art': 3,
                          'Computer Science and Engineering': 3,
                          'Construction Engineering': 3,
                          'Cost and Management Accountant': 3,
                          'CS- Company Secretary': 3,
                          'Diesel Mechanical Engineering': 3,
                          'Diploma in Art and Craft': 2,
                          'Diploma in Ayurvedic Pharmacy': 2,
                          'Diploma in Civil Engineering': 2,
                          'Diploma in Clinical Research and Pharmacovigilance': 2,
                          'Diploma in Computer Engineering': 2,
                          'Diploma in Drug Store Management': 2,
                          'Diploma in Electrical and Electronics Engineering': 2,
                          'Diploma in Electrical Engineering': 2,
                          'Diploma in Electronics and Communication Engineering': 2,
                          'Diploma in Fine Arts': 2,
                          'Diploma in Information Technology': 2,
                          'Diploma in Mechanical Engineering': 2,
                          'Diploma in Performing Arts': 2,
                          'Diploma in Pharmaceutical Marketing': 2,
                          'Diploma in Pharmacy': 2,
                          'Doctor of Pharmacy': 2,
                          'Draughtsman (Civil) Engineering': 2,
                          'Draughtsman (Mechanical) Engineering': 2,
                          'Dress Making': 2,
                          'E-Commerce': 2,
                          'Electrical and Electronics Engineering': 3,
                          'Electronics and Communication Engineering': 3,
                          'Fitter Engineering': 3,
                          'Foundry Man Engineering': 3,
                          'Fruits and Vegetable Processing': 2,
                          'Graphic Designing': 3,
                          'Hair and Skin Care': 2,
                          'Hand Compositor': 2,
                          'Industry Certificate in Biopharmaceutical Technology': 2,
                          'Industry Certificate in Pharma Product Management': 2,
                          'Industry Certificate in Pharma Sales and Marketing Management': 2,
                          'Industry Certificate in Pharmaceutical Chemistry': 2,
                          'Industry Certificate in Pharmaceutical Formulation and Entrepreneurship': 2,
                          'Industry Certificate in Pharmaceutical Packaging': 2,
                          'Industry Certificate in Pharmaceutical Process Engineering': 2,
                          'Industry Certificate in Pharmaceutical Technology Transfer': 2,
                          'Industry Certificate in Pharmacovigilance': 2,
                          'Information Technology and ESM Engineering': 3,
                          'Instrumentation Engineering': 3,
                          'Integarted Law Program- BBA LL.B': 3,
                          'Integrated Law Course- BA + LL.B': 3,
                          'Integrated Law Program- B.Com LL.B.': 3,
                          'Intermediate': 2,
                          'Jute Production Technology': 3,
                          'Leather Goods Maker': 2,
                          'Letter Press Machine Mender': 2,
                          'M Pharm (Biopharmaceutics & Pharmacokinetics)': 4,
                          'M Pharm (Biopharmaceutics)': 4,
                          'M Pharm (Biotechnology)': 4,
                          'M Pharm (Clinical Pharmacy)': 4,
                          'M Pharm (Cosmeceuticals)': 4,
                          'M Pharm (Drug Discovery and Drug Development)': 4,
                          'M Pharm (Industrial Pharmacy)': 4,
                          'M Pharm (Medicinal Chemistry)': 4,
                          'M Pharm (Pharmaceutical Chemistry)': 4,
                          'M Pharm (Pharmaceutical Market and Management)': 4,
                          'M Pharm (Pharmaceutical Quality Assurance)': 4,
                          'M Pharm (Pharmaceutics)': 4,
                          'M Pharmacy (Master of Pharmacy)': 4,
                          'Machinist Engineering': 4,
                          'Manufacture Foot Wear': 4,
                          'Master of Fine Arts (MFA)': 4,
                          'Master of Performing Arts (MPA)': 4,
                          'Master of Visual Arts (MVA)': 4,
                          'Masters of Science': 4,
                          'MBA Agriculture & Food Business': 4,
                          'Mech. Instrument Engineering': 3,
                          'Mechanic Electronics Engineering': 3,
                          'Mechanic Radio & TV Engineering': 3,
                          'Motor Diver Cum Mechanic Engineering': 3,
                          'MSc (Clinical Research and Pharmacovigilance)': 4,
                          'MSc (Pharmacology)': 4,
                          'MSc Agriculture': 4,
                          'MSc in Agricultural Economics': 4,
                          'MSc in Agriculture': 4,
                          'MSc in Agronomy': 4,
                          'MSc in Animal Husbandry': 4,
                          'MSc in Plant Biotechnology': 4,
                          'MSc in Plant Pathology': 4,
                          'MSc in Plant Physiology': 4,
                          'MSc in Seed Science and Technology': 4,
                          'MSc in Sericulture': 4,
                          'MSc in Soil Science': 4,
                          'MTech Agricultural Engineering': 4,
                          'MTech Agricultural Water Management': 4,
                          'MTech Agriculture Process & Food Engineering': 4,
                          'MTech Agriculture Systems, and Management': 4,
                          'MTech/ME in Agriculture Biotechnology, ': 4,
                          'Oyster Mushroom Production Technology': 3,
                          'Petroleum Engineering': 3,
                          'PG Diploma in Drug Store Management': 4,
                          'PG Diploma in Herbal Products': 4,
                          'PG Diploma in Pharmaceutical Chemistry': 4,
                          'PG Diploma in Pharmaceutical Management': 4,
                          'PG Diploma in Pharmaceutical Quality Assurance': 4,
                          'PG Diploma in Principles of Clinical Pharmacology': 4,
                          'PGD in Agricultural Extension Management': 4,
                          'PGD in Art and Design': 4,
                          'PGD in Cinematography': 4,
                          'PGD in Management: Agribusiness and Plantation Management': 4,
                          'PGD in Museology': 4,
                          'PGD in Technology Management in Agriculture': 4,
                          'PGD in TV Direction': 4,
                          'Ph.D. (Pharmaceutical Biotechnology)': 5,
                          'Ph.D. (Pharmaceutical Chemistry)': 5,
                          'Ph.D. (Pharmaceutical Medicine)': 5,
                          'Ph.D. (Pharmaceutical Sciences)': 5,
                          'Ph.D. (Pharmacognosy & Phytochemistry)': 5,
                          'Ph.D. (Pharmacognosy)': 5,
                          'Ph.D. (Pharmacology)': 5,
                          'Ph.D. (Pharmacy Practice)': 5,
                          'Ph.D. (Phytopharmacy & Phytomedicine)': 5,
                          'PhD': 5,
                          'PhD in Agriculture Economy': 5,
                          'PhD in Agriculture, Horticulture': 5,
                          'PhD in Agronomy': 5,
                          'PhD in Arts': 5,
                          'PhD in Genetics & Plant Breeding': 5,
                          'PhD in Plant Pathology': 5,
                          'Physiotherapy': 2,
                          'Plant Protection': 2,
                          'Power Engineering': 3,
                          'Professional Certificate In Pharmacovigilance': 2,
                          'Pump Operator Engineering': 3,
                          'Refrigeration Engineering': 3,
                          'Robotics Engineering': 3,
                          'Safety and Pharmacovigilance Certificate Program': 2,
                          'Secretarial Practice': 2,
                          'Sheet Metal Worker Engineering': 3,
                          'Smart Manufacturing & Automation': 3,
                          'Structural Engineering': 3,
                          'Surveyor Engineering': 3,
                          'Tally ERP Course': 2,
                          'Textile Engineering': 3,
                          'Tool and Die Maker Engineering': 3,
                          'Transportation Engineering': 3,
                          'Turner Engineering': 3,
                          'Uneducated': 0,
                          'Water Management For Crop Production': 2,
                          'Web Designing & Development': 3
                                }
            if x in education_mapping:
                qualification_encoded = education_mapping[x]
            else:
                qualification_encoded = 0
                                  
    except ValueError: 
             qualification_encoded=0
                              
    return qualification_encoded    


def getedu_lmt(qualification):
    if "All Education levels" in qualification or "above" in qualification:
        edu_lmt="gte"
    elif "below" in qualification or "Graduate" in qualification:
        edu_lmt="lte"
    else:
        edu_lmt="lte" 
    return edu_lmt

#Function to read doc format resume
def readpdf(file):
    doc = fitz.Document(stream=file, filetype="pdf")
    text = ""
    for i in range(doc.pageCount):
        page = doc.loadPage(i)
        pagetext = page.get_text("text").strip()
        text += (pagetext+ '\n')
    return text

def iter_block_items(parent):
    """
    Yield each paragraph and table child within *parent*, in document order.
    Each returned value is an instance of either Table or Paragraph.
    """
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

 
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def iter_unique_cells(row, prior_tc):
    for cell in row.cells:
        this_tc = cell._tc
        if this_tc in prior_tc:
            continue
        prior_tc.append(this_tc)
        yield cell
        
def get_text(filename):
    doc = docx.Document(filename)
    fulltext = []
    for para in doc.paragraphs:
        fulltext.append(para.text)
    return '\n'.join(fulltext)

def readdoc(file):
    doc = docx.Document(file)
    text = ''
    for block in iter_block_items(doc):
        if isinstance(block,docx.table.Table):
            prior_tc = []
            for row in block.rows:
                for cell in iter_unique_cells(row, prior_tc):
                    text += (cell.text + '\t')
                text += '\n'
            continue
        para = block
        if str(para.text).strip() == '':
            continue
        text += (para.text + '\n')
    return text

#Function for creating dataFrame from json
def user_json_to_df(json_object):
    user_data = json.loads(json_object)
    df = pd.json_normalize(user_data['results'])
    return df

#Function for creating similarity_match
def similarity_match(x):  
    val = ','.join(str(v) for v in x)
    cleanval = clean_text(val)
    # # # Generating Sentence Bert
    embed_user = convert_sentancevec(cleanval)
    vector=convert_array(embed_user)
    return vector

#Function for finding distance between user and jobposting
def get_distance(x,y,user_lat,user_long):
    try:
        distance1 = (distance.distance((x,y),(user_lat,user_long))) 
    
    except AttributeError:
                    distance1 = 0
                    
    except NameError:
                    distance1 = 0
                    
    except KeyError:
                    distance1 = 0
    
    return distance1



 ##### Function for covert latitude and longitude"

def get_latlong(x):
        try:
                coordinates1=geocoder.geocode(x, timeout=None)                
                latitude,longitude=(coordinates1.latitude),(coordinates1.longitude)
               
    
        except AttributeError:
                    latitude = "0"
                    longitude = "0"
                    
        except NameError:
                    latitude = "0"
                    longitude = "0"
                    
        except KeyError:
                    latitude = "0"
                    longitude = "0"
    
        return [latitude,longitude]   

def listto_string(inputlist):
    outputstring=[]
    for each in inputlist:
        for values in each.values():
            outputstring.append(values)
    outputstring = ' '.join(outputstring).split() 
    
    return outputstring    
 
def inputto_list(input):
    outputlist=[]
    for each in input:
        for values in each.values():
            outputlist.append(values)
    return outputlist 
